''' 
* ======================================================================================================= *
*    Script Description:                                                                                  *
*    Python Script to combine two dataframes, add dataframe name on top of the respective dataframe,      *
*    and save the resulting dataframe in a word file.                                                     *
*                                                                                                         *
*    Date Created: 18 December 2020                                                                       *  
*                                                                                                         *
*    Dependencies:                                                                                        *
*    Following package(s) must be installed before running the script:                                    *
*          ==> python-docx                                                                                *
*    Command to install package(s):                                                                       *
*    pip3 install python-docx                                                                             *
*                                                                                                         *
*    Instructions to run the Script:                                                                      *
*    Open Terminal(MacOS or Linux) / Command Prompt(Windows) in the directory where the script is present *
*    Run the following command to create the outputdoc file in current directory:                         *
*    python3 <script_name>.py                                                                             *                       
*                                                                                                         *                    
*    To create the output file in any other location, run the below command:                              *                    
*    python3 <script_name>.py -op <destination directory path>                                            *                               
*    Example: python3 DF_to_Word.py -op /Users/dkotnala/Desktop                                           *   
*                                                                                                         *
*    Developed By:                                                                                        *
*        ____                              __                  __ __      __              __              * 
*       / __ \___  ___  ____  ____ _____  / /______ ______    / //_/___  / /_____  ____ _/ /___ _         *
*      / / / / _ \/ _ \/ __ \/ __ `/ __ \/ //_/ __ `/ ___/   / ,< / __ \/ __/ __ \/ __ `/ / __ `/         *
*     / /_/ /  __/  __/ /_/ / /_/ / / / / ,< / /_/ / /      / /| / /_/ / /_/ / / / /_/ / / /_/ /          *
*    /_____/\___/\___/ .___/\__,_/_/ /_/_/|_|\__,_/_/      /_/ |_\____/\__/_/ /_/\__,_/_/\__,_/           *
*                   /_/                                                                                   *
*                                                                                                         *
* ======================================================================================================= * 
'''

# Importing the required libraries
import os
import docx
from docx import Document
import pandas as pd
import argparse

# Find the current directory
dir_path = os.path.dirname(os.path.realpath(__file__))

# Argument Parser for making the script user-interactive. 
# Used for taking destination directory for output file (Word Doc)
parser = argparse.ArgumentParser()
#parser.add_argument("--input", help = "input filename")
parser.add_argument("-op", "--output", default = dir_path, 
                                       help = "Please provide the path of destination directory for output file. \
                                               Default directory is current path where the script is present.")
args = parser.parse_args()
dest_dir = args.output

# Hard-coding dataframe for test -- Dataframe 1
df1 = pd.DataFrame({'a1': ['A0', 'A1', 'A2', 'A3'],
                    'a2': ['B0', 'B1', 'B2', ''],
                    'a3':[1,2,3,4],
                    'a4': ['B0', 'B1', 'B2', ''],
                    'a5':[1,2,3,4],
                    'a6':['z','e','e','k']})

# Hard-coding dataframe for test -- Dataframe 2
df2 = pd.DataFrame({'b1': ['0', '4', '9', '3'],
                    'b2': ['1', '2', '3', '4'],
                    'b3': ['5', '6', '7', '8'],
                    'b4': ['2', '3', '4', '6'],
                    'b5': ['0', '', '9', '3'],
                    'b6': ['1', '2', '3', '4'],
                    'b7': ['5', '6', '', '8'],
                    'b8': ['2', '3', '4', '6']})

try:
    # Concatenating the two dataframes
    concatenated_df = pd.concat([df1, df2], axis=1)
    len_df1 = int(len(df1.columns))
    len_df2 = int(len(df2.columns))
    mid_df1 = int(len(df1.columns)/2)
    mid_df2 = int(len(df2.columns)/2)
    col_list = []
    for i in concatenated_df.columns:
        col_list.append(i)

    # Creating a table structure in the doc as per our Data Frame
    doc=Document()
    run = doc.add_paragraph().add_run()
    font = run.font
    from docx.shared import Pt
    font.size = Pt(12)
    doc.add_heading("Combined Dataframe", level=1)
    table = doc.add_table(rows = concatenated_df.shape[0]+2, cols = concatenated_df.shape[1])
    table.style = 'Table Grid'

    #start_of_df2 = int(len(df1.columns)) + mid_df2

    # Merging cells for DataFrame 1 header
    for i in range(pd.Series(col_list).shape[0]):
        if i < (len_df1-1):
            A = table.cell(0,i)
            A = A.merge(table.cell(0, i+1))

    # Center alignment of text fields in the table (DF in word) header and shading the cells
    table.cell(0,1).paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    shading_elm = docx.oxml.parse_xml(r'<w:shd {} w:fill="FFD2A2"/>'.format(docx.oxml.ns.nsdecls('w')))
    table.cell(0,1)._tc.get_or_add_tcPr().append(shading_elm)

    # Merging cells for DataFrame 2 header
    for i in range(pd.Series(col_list).shape[0] - len_df1 - 1):
        b = table.cell(0,(len_df1 + i))
        b = b.merge(table.cell(0, (len_df1 + i+1)))

    # For Column Names 
    for i in range(pd.Series(col_list).shape[0]):
        # Making the headers BOLD 
        run = table.cell(1, i).paragraphs[0].add_run(str(col_list[i]))
        run.font.bold = True
        # Font color for column names
        # run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)  
        table.cell(1,i).paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        # Different shading for columns of different DFs
        if i < len_df1:
            shading_elm = docx.oxml.parse_xml(r'<w:shd {} w:fill="D7EEFF"/>'.format(docx.oxml.ns.nsdecls('w')))
            table.cell(1,i)._tc.get_or_add_tcPr().append(shading_elm)
        else:
            shading_elm = docx.oxml.parse_xml(r'<w:shd {} w:fill="D3FBD2"/>'.format(docx.oxml.ns.nsdecls('w')))
            table.cell(1,i)._tc.get_or_add_tcPr().append(shading_elm)

    # For putting values of Dataframe in the doc (table in doc)
    for i in range(concatenated_df.shape[0]):
        for j in range(concatenated_df.shape[1]):
            table.cell(2+i,j).text=str(concatenated_df.values[i,j])
            table.cell(i+2,j).paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    # Putting Header name for Dataframes in the merged header cells
    run_dfname = table.cell(0, 1).paragraphs[0].add_run("Dataframe1")
    run_dfname.font.bold = True
    run_dfname = table.cell(0, (len_df1 + mid_df2)).paragraphs[0].add_run("Dataframe2")
    run_dfname.font.bold = True
    shading_elm = docx.oxml.parse_xml(r'<w:shd {} w:fill="D6D6D6"/>'.format(docx.oxml.ns.nsdecls('w')))

    table.cell(0,(len_df1 + mid_df2))._tc.get_or_add_tcPr().append(shading_elm)
    table.cell(0,(len_df1 + mid_df2)).paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

except Exception as e:
    print("Exception occurred while merging the two DFs. \n", e)

# Hard-coding the path in test code
#doc.save('/Users/dkotnala/mywork/demo.docx') # Save document

# Saving the document in specified directory. 
# If directory is not provided while running the script, the default directory will be taken as current directory.
try:
    doc.save(os.path.join(dest_dir, "Combined_DF_Doc.docx"))
    print("File saved: ", dest_dir)
except Exception as e:
    print("\nException occurred while saving the file. \nPlease check the destination directory path provided.\n", e)



# * =================================== E N D  O F  S C R I P T ============================================ *

